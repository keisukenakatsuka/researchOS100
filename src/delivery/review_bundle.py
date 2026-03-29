# src/delivery/review_bundle.py
"""Review Bundle Generator — assemble research artifacts into a single .docx.

Reads paper_draft.md, review_report.md, and evidence_sufficiency.json from
a run directory and produces a review_bundle.docx suitable for Kindle reading.

Usage::

    from src.delivery.review_bundle import build_review_bundle

    docx_path = build_review_bundle(run_dir)
"""

from __future__ import annotations

import glob
import json
import logging
import re
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Dict, List, Optional

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

logger = logging.getLogger(__name__)


# ------------------------------------------------------------------
# Public API
# ------------------------------------------------------------------

def build_review_bundle(
    run_dir: Path,
    *,
    output_dir: Optional[Path] = None,
) -> Path:
    """Build review_bundle.docx from run artifacts.

    Args:
        run_dir: Path to run directory (data/lit_review/{run_id}/)
        output_dir: Override output directory (default: run_dir)

    Returns:
        Path to generated .docx file
    """
    out = output_dir or run_dir
    run_id = run_dir.name
    now = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")

    doc = Document()
    _set_default_font(doc, "Times New Roman", Pt(11))

    # --- Title page ---
    rq_title = _load_rq_title(run_dir)
    _add_title_page(doc, run_id=run_id, rq_title=rq_title, generated_at=now)

    # --- Section 1: Paper Draft ---
    doc.add_page_break()
    _add_section_heading(doc, "1. Paper Draft")
    paper_md = _read_file(run_dir / "paper_draft.md")
    if paper_md:
        _render_markdown(doc, paper_md)
    else:
        _add_fallback(doc, "paper_draft.md not found")

    # --- Section 2: Review Report ---
    doc.add_page_break()
    _add_section_heading(doc, "2. Review Report")
    review_md = _read_file(run_dir / "review_report.md")
    if review_md:
        _render_markdown(doc, review_md)
    else:
        _add_fallback(doc, "review_report.md not found")

    # --- Section 3: Evidence Sufficiency ---
    doc.add_page_break()
    _add_section_heading(doc, "3. Evidence Sufficiency")
    sufficiency_files = _find_sufficiency_files(run_dir)
    if sufficiency_files:
        for i, (hyp_id, data) in enumerate(sufficiency_files, 1):
            _render_sufficiency(doc, hyp_id, data, index=i)
    else:
        _add_fallback(doc, "No evidence_sufficiency.json files found")

    # --- Save ---
    ts = datetime.now().strftime("%Y%m%d_%H%M")
    filename = f"review_bundle_{ts}.docx"
    out_path = out / filename
    doc.save(str(out_path))
    logger.info("Saved review_bundle: %s", out_path)
    return out_path


# ------------------------------------------------------------------
# Title page
# ------------------------------------------------------------------

def _add_title_page(doc: Document, *, run_id: str, rq_title: str, generated_at: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Review Bundle")
    run.bold = True
    run.font.size = Pt(24)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.add_run(f"Run: {run_id}").font.size = Pt(12)

    if rq_title:
        p3 = doc.add_paragraph()
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r3 = p3.add_run(rq_title)
        r3.font.size = Pt(11)
        r3.italic = True

    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p4.add_run(f"Generated: {generated_at}").font.size = Pt(10)

    # Table of Contents
    doc.add_paragraph()
    toc = doc.add_paragraph()
    toc.alignment = WD_ALIGN_PARAGRAPH.LEFT
    toc_run = toc.add_run("Contents")
    toc_run.bold = True
    toc_run.font.size = Pt(14)

    for item in ["1. Paper Draft", "2. Review Report", "3. Evidence Sufficiency"]:
        p = doc.add_paragraph(item, style="List Number")


# ------------------------------------------------------------------
# Markdown → docx renderer
# ------------------------------------------------------------------

def _render_markdown(doc: Document, text: str):
    """Convert Markdown text to docx paragraphs.

    Supported patterns:
    - # / ## / ### / #### headings
    - **bold** and *italic*
    - - bullet lists
    - 1. numbered lists
    - | pipe tables
    - ``` code blocks
    - --- horizontal rule
    """
    lines = text.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]

        # Code block
        if line.strip().startswith("```"):
            i += 1
            code_lines = []
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            i += 1  # skip closing ```
            _add_code_block(doc, "\n".join(code_lines))
            continue

        # Horizontal rule
        if line.strip() in ("---", "***", "___"):
            _add_horizontal_rule(doc)
            i += 1
            continue

        # Headings
        heading_match = re.match(r"^(#{1,4})\s+(.*)", line)
        if heading_match:
            level = len(heading_match.group(1))
            text_content = heading_match.group(2).strip()
            doc.add_heading(text_content, level=level)
            i += 1
            continue

        # Table (consecutive pipe-delimited lines)
        if "|" in line and line.strip().startswith("|"):
            table_lines = []
            while i < len(lines) and "|" in lines[i] and lines[i].strip().startswith("|"):
                table_lines.append(lines[i])
                i += 1
            _add_table(doc, table_lines)
            continue

        # Bullet list
        if re.match(r"^[-*]\s+", line.strip()):
            text_content = re.sub(r"^[-*]\s+", "", line.strip())
            p = doc.add_paragraph(style="List Bullet")
            _add_formatted_runs(p, text_content)
            i += 1
            continue

        # Numbered list
        if re.match(r"^\d+\.\s+", line.strip()):
            text_content = re.sub(r"^\d+\.\s+", "", line.strip())
            p = doc.add_paragraph(style="List Number")
            _add_formatted_runs(p, text_content)
            i += 1
            continue

        # Empty line
        if not line.strip():
            i += 1
            continue

        # Regular paragraph
        p = doc.add_paragraph()
        _add_formatted_runs(p, line)
        i += 1


def _add_formatted_runs(paragraph, text: str):
    """Parse inline formatting (**bold**, *italic*) and add as runs."""
    # Pattern: **bold** or *italic* (non-greedy)
    pattern = r"(\*\*(.+?)\*\*|\*(.+?)\*)"
    last_end = 0

    for m in re.finditer(pattern, text):
        # Add text before match
        if m.start() > last_end:
            paragraph.add_run(text[last_end:m.start()])

        if m.group(2):  # **bold**
            run = paragraph.add_run(m.group(2))
            run.bold = True
        elif m.group(3):  # *italic*
            run = paragraph.add_run(m.group(3))
            run.italic = True

        last_end = m.end()

    # Add remaining text
    if last_end < len(text):
        paragraph.add_run(text[last_end:])


def _add_table(doc: Document, lines: List[str]):
    """Parse markdown pipe table and add as docx table."""
    rows = []
    for line in lines:
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        # Skip separator row (e.g., |---|---|)
        if all(re.match(r"^[-:]+$", c) for c in cells):
            continue
        rows.append(cells)

    if not rows:
        return

    n_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=n_cols)
    table.style = "Table Grid"

    for r_idx, row in enumerate(rows):
        for c_idx, cell in enumerate(row):
            if c_idx < n_cols:
                table_cell = table.cell(r_idx, c_idx)
                table_cell.text = ""
                p = table_cell.paragraphs[0]
                _add_formatted_runs(p, cell)
                # Bold header row
                if r_idx == 0:
                    for run in p.runs:
                        run.bold = True


def _add_code_block(doc: Document, code: str):
    """Add code block as monospace paragraph."""
    p = doc.add_paragraph()
    run = p.add_run(code)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)


def _add_horizontal_rule(doc: Document):
    """Add a visual separator."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("" + "\u2500" * 40 + "")
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    run.font.size = Pt(8)


# ------------------------------------------------------------------
# Evidence Sufficiency rendering
# ------------------------------------------------------------------

def _render_sufficiency(doc: Document, hyp_id: str, data: Dict[str, Any], index: int):
    """Render a single hypothesis sufficiency result."""
    stmt = data.get("hypothesis_statement", "")
    if not stmt:
        stmt = hyp_id

    # Heading
    display_stmt = stmt[:80] + "..." if len(stmt) > 80 else stmt
    doc.add_heading(f"3.{index} {hyp_id[:20]}", level=3)

    # Statement
    p = doc.add_paragraph()
    p.add_run("Statement: ").bold = True
    p.add_run(display_stmt).italic = True

    # Status table
    result = data.get("result", "unknown")
    recommendation = data.get("recommendation", "unknown")
    consensus = data.get("consensus_support", {})
    max_strength = consensus.get("max_strength", "unknown")
    support_count = consensus.get("count", 0)
    gaps = data.get("gap_concerns", {})
    gap_count = gaps.get("count", 0)

    table = doc.add_table(rows=5, cols=2)
    table.style = "Table Grid"
    _set_table_row(table, 0, "Verdict", result.upper())
    _set_table_row(table, 1, "Recommendation", recommendation)
    _set_table_row(table, 2, "Consensus Strength", max_strength)
    _set_table_row(table, 3, "Support Count", str(support_count))
    _set_table_row(table, 4, "Critical Gaps", str(gap_count))

    # Key evidence (consensus findings)
    findings = consensus.get("relevant_findings", [])
    if findings:
        doc.add_heading("Key Evidence", level=4)
        for f in findings:
            doc.add_paragraph(f, style="List Bullet")

    # Critical gaps
    critical = gaps.get("critical_gaps", [])
    if critical:
        doc.add_heading("Critical Gaps", level=4)
        for g in critical:
            doc.add_paragraph(g, style="List Bullet")

    # Suggested queries
    queries = data.get("suggested_queries", [])
    if queries:
        doc.add_heading("Suggested Queries", level=4)
        for q in queries:
            doc.add_paragraph(q, style="List Bullet")

    doc.add_paragraph()  # spacing


def _set_table_row(table, row_idx: int, label: str, value: str):
    table.cell(row_idx, 0).text = label
    for run in table.cell(row_idx, 0).paragraphs[0].runs:
        run.bold = True
    table.cell(row_idx, 1).text = value


# ------------------------------------------------------------------
# Helpers
# ------------------------------------------------------------------

def _set_default_font(doc: Document, name: str, size):
    """Set default font for the document."""
    style = doc.styles["Normal"]
    font = style.font
    font.name = name
    font.size = size


def _add_section_heading(doc: Document, text: str):
    """Add a top-level section heading."""
    doc.add_heading(text, level=1)


def _add_fallback(doc: Document, message: str):
    """Add a fallback paragraph when content is unavailable."""
    p = doc.add_paragraph()
    run = p.add_run(f"[{message}]")
    run.italic = True
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)


def _read_file(path: Path) -> Optional[str]:
    """Read file content, returning None if missing or empty."""
    if not path.exists():
        logger.warning("File not found: %s", path)
        return None
    text = path.read_text(encoding="utf-8")
    if not text.strip():
        logger.warning("File is empty: %s", path)
        return None
    return text


def _load_rq_title(run_dir: Path) -> str:
    """Load RQ title from rq_context.json or export_bundle.json."""
    for name in ("rq_context.json", "export_bundle.json"):
        path = run_dir / name
        if path.exists():
            try:
                data = json.loads(path.read_text())
                title = data.get("rq_title", data.get("title", ""))
                if title:
                    return title
            except (json.JSONDecodeError, OSError):
                pass
    return ""


def _find_sufficiency_files(run_dir: Path) -> List[tuple]:
    """Find all evidence_sufficiency.json files under hyp_literature/.

    Returns list of (hypothesis_id, data_dict) sorted by hypothesis_id.
    """
    pattern = str(run_dir / "hyp_literature" / "*" / "evidence_sufficiency.json")
    results = []
    for path in sorted(glob.glob(pattern)):
        try:
            data = json.loads(Path(path).read_text())
            hyp_id = data.get("hypothesis_id", Path(path).parent.name)
            results.append((hyp_id, data))
        except (json.JSONDecodeError, OSError) as e:
            logger.warning("Failed to load %s: %s", path, e)
    return results
