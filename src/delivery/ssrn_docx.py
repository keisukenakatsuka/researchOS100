# src/delivery/ssrn_docx.py
"""SSRN submission docx generator.

Converts a paper markdown draft to a Word document suitable for SSRN upload.

Key enhancements over review_bundle.py:
- Escaped significance stars (\\*\\*\\*) render as literal *** in tables
- Japanese literature corpus is stripped from the output
- Academic paper title page (title, author, date)
- Horizontal rules (---) are skipped rather than rendered
- Table font size reduced to 10pt for readability
"""

from __future__ import annotations

import logging
import re
from datetime import datetime
from pathlib import Path
from typing import List, Optional, Tuple

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

logger = logging.getLogger(__name__)

# Placeholder for markdown-escaped asterisks (\*) so they survive
# bold/italic parsing and are restored as literal * in the final output.
_STAR = "\ufffd"


# ------------------------------------------------------------------
# Public API
# ------------------------------------------------------------------

def build_ssrn_docx(
    md_path: Path,
    *,
    output_dir: Optional[Path] = None,
    author: str = "",
    date: str = "",
) -> Path:
    """Convert paper markdown to SSRN-ready docx.

    Args:
        md_path: Path to the paper markdown file.
        output_dir: Where to write the docx (default: same dir as md_path).
        author: Author name for the title page.
        date: Paper date string for the title page.

    Returns:
        Path to the generated .docx file.
    """
    out = output_dir or md_path.parent
    text = md_path.read_text(encoding="utf-8")

    # Strip Japanese literature corpus (after English appendix tables)
    text = _strip_literature_corpus(text)

    # Pre-process: convert \* → placeholder so bold/italic parsing ignores them
    text = text.replace("\\*", _STAR)

    doc = Document()
    _setup_styles(doc)

    lines = text.split("\n")

    # Extract title from first # heading
    title, body_start = _extract_title(lines)

    # Build document
    _add_title_page(doc, title=title, author=author, date=date)
    _render_body(doc, lines[body_start:])

    # Save
    ts = datetime.now().strftime("%Y%m%d_%H%M")
    out_path = out / f"paper_ssrn_{ts}.docx"
    doc.save(str(out_path))
    logger.info("SSRN docx saved: %s", out_path)
    return out_path


# ------------------------------------------------------------------
# Pre-processing
# ------------------------------------------------------------------

def _strip_literature_corpus(text: str) -> str:
    """Remove the Japanese literature corpus section.

    Cuts at '## Literature Corpus' (which begins the non-English appendix)
    while preserving the English appendix tables above it.
    """
    marker = "## Literature Corpus"
    idx = text.find(marker)
    if idx > 0:
        before = text[:idx].rstrip()
        # Also strip the preceding --- divider if present
        if before.endswith("---"):
            before = before[:-3].rstrip()
        return before
    return text


def _extract_title(lines: List[str]) -> Tuple[str, int]:
    """Find and extract the first # heading as the paper title."""
    for i, line in enumerate(lines):
        m = re.match(r"^#\s+(.*)", line)
        if m:
            return _restore_stars(m.group(1).strip()), i + 1
    return "", 0


# ------------------------------------------------------------------
# Document setup
# ------------------------------------------------------------------

def _setup_styles(doc: Document):
    """Set default font and spacing for the document."""
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.space_before = Pt(0)


def _add_title_page(
    doc: Document,
    *,
    title: str,
    author: str,
    date: str,
):
    """Add a centered title block (title, author, date)."""
    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = "Times New Roman"

    # Author
    if author:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(author)
        r.font.size = Pt(12)
        r.font.name = "Times New Roman"

    # Date
    if date:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(date)
        r.font.size = Pt(11)
        r.font.name = "Times New Roman"

    doc.add_paragraph()  # spacing before body


# ------------------------------------------------------------------
# Body renderer
# ------------------------------------------------------------------

def _render_body(doc: Document, lines: List[str]):
    """Line-by-line markdown → docx conversion."""
    i = 0
    while i < len(lines):
        line = lines[i]

        # --- Code block ---
        if line.strip().startswith("```"):
            i += 1
            code_lines: List[str] = []
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            if i < len(lines):
                i += 1  # skip closing ```
            _add_code_block(doc, "\n".join(code_lines))
            continue

        # --- Horizontal rule → skip ---
        stripped = line.strip()
        if stripped in ("---", "***", "___"):
            i += 1
            continue

        # --- Headings ---
        m = re.match(r"^(#{1,4})\s+(.*)", line)
        if m:
            level = len(m.group(1))
            heading_text = _restore_stars(m.group(2).strip())
            doc.add_heading(heading_text, level=level)
            i += 1
            continue

        # --- Pipe table ---
        if stripped.startswith("|") and "|" in stripped[1:]:
            table_lines: List[str] = []
            while (
                i < len(lines)
                and lines[i].strip().startswith("|")
                and "|" in lines[i].strip()[1:]
            ):
                table_lines.append(lines[i])
                i += 1
            _add_table(doc, table_lines)
            continue

        # --- Bullet list ---
        bullet_m = re.match(r"^[-]\s+(.*)", stripped)
        if bullet_m:
            p = doc.add_paragraph(style="List Bullet")
            _add_formatted_runs(p, bullet_m.group(1))
            i += 1
            continue

        # --- Numbered list ---
        num_m = re.match(r"^\d+\.\s+(.*)", stripped)
        if num_m:
            p = doc.add_paragraph(style="List Number")
            _add_formatted_runs(p, num_m.group(1))
            i += 1
            continue

        # --- Empty line → skip ---
        if not stripped:
            i += 1
            continue

        # --- Regular paragraph ---
        p = doc.add_paragraph()
        _add_formatted_runs(p, line)
        i += 1


# ------------------------------------------------------------------
# Inline formatting
# ------------------------------------------------------------------

def _add_formatted_runs(paragraph, text: str):
    """Parse **bold** and *italic* inline, restoring star placeholders."""
    # Match **bold** first, then *italic* (non-greedy)
    pattern = r"(\*\*(.+?)\*\*|\*(.+?)\*)"
    last_end = 0

    for m in re.finditer(pattern, text):
        # Plain text before the match
        if m.start() > last_end:
            paragraph.add_run(_restore_stars(text[last_end : m.start()]))

        if m.group(2):  # **bold**
            run = paragraph.add_run(_restore_stars(m.group(2)))
            run.bold = True
        elif m.group(3):  # *italic*
            run = paragraph.add_run(_restore_stars(m.group(3)))
            run.italic = True

        last_end = m.end()

    # Remaining text after last match
    if last_end < len(text):
        paragraph.add_run(_restore_stars(text[last_end:]))


# ------------------------------------------------------------------
# Tables
# ------------------------------------------------------------------

def _add_table(doc: Document, lines: List[str]):
    """Convert markdown pipe table to a Word table."""
    rows: List[List[str]] = []
    for line in lines:
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        # Skip separator rows (|---|---|)
        if all(re.match(r"^[-:]+$", c) for c in cells if c):
            continue
        rows.append(cells)

    if not rows:
        return

    n_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=n_cols)
    table.style = "Table Grid"

    for r_idx, row in enumerate(rows):
        for c_idx, cell_text in enumerate(row):
            if c_idx >= n_cols:
                continue
            tc = table.cell(r_idx, c_idx)
            tc.text = ""
            p = tc.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.space_before = Pt(0)
            _add_formatted_runs(p, cell_text)
            # Bold the header row
            if r_idx == 0:
                for run in p.runs:
                    run.bold = True
            # Smaller font for table cells
            for run in p.runs:
                run.font.size = Pt(10)
                run.font.name = "Times New Roman"

    # Spacing after table
    doc.add_paragraph()


# ------------------------------------------------------------------
# Code blocks
# ------------------------------------------------------------------

def _add_code_block(doc: Document, code: str):
    """Render code block as indented monospace text (for equations etc.)."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(36)
    run = p.add_run(_restore_stars(code))
    run.font.name = "Courier New"
    run.font.size = Pt(10)


# ------------------------------------------------------------------
# Helpers
# ------------------------------------------------------------------

def _restore_stars(text: str) -> str:
    """Convert placeholders back to literal asterisks."""
    return text.replace(_STAR, "*")
